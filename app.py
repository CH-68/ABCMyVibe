import io
import streamlit as st
from pypdf import PdfReader
from helper_functions.utility import check_password
from src.crew import setup_and_run_crew
from docx import Document

st.set_page_config(page_title="Compliance Verifier", page_icon="🛡️")
st.title("Agentic Compliance Verifier")
st.caption(
    "CrewAI-based verification for repeatable, auditable policy checks to aid human-in-the-loop."
)

if "report" not in st.session_state:
    st.session_state.report = None

# Do not continue if check_password is not True.
if not check_password():
        st.stop()

def extract_pdf_text(uploaded_file) -> str:
    if uploaded_file is None:
        return ""

    reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n".join(pages).strip()


with st.sidebar:
    st.header("Policy and target documents")
    policy_file = st.file_uploader("Upload policy PDF", type=["pdf"], key="policy_pdf")
    target_file = st.file_uploader("Upload target document PDF", type=["pdf"], key="target_pdf")

    st.markdown("---")
    if st.button("Run verification", disabled=not (policy_file and target_file)):
        with st.spinner("Running verification..."):
            policy_text = extract_pdf_text(policy_file)
            target_text = extract_pdf_text(target_file)
            if not policy_text or not target_text:
                st.error("Unable to extract readable text from one or both PDFs.")
            else:
                st.session_state.report = setup_and_run_crew(
                    policy_text=policy_text,
                    document_text=target_text,
                )
                st.success("Verification report generated.")


if st.session_state.report:
    report = st.session_state.report
    summary = report.get("summary", {})
    counts = summary.get("counts", {})

    st.subheader("Verification summary")
    pass_count = counts.get("pass", 0)
    fail_count = counts.get("fail", 0)
    insuff_count = counts.get("insufficient evidence", 0)
    col_pass, col_fail, col_review = st.columns(3)
    col_pass.metric("Passed", pass_count, "✔️")
    col_fail.metric("Failed", fail_count, "❌")
    col_review.metric("Further review", insuff_count, "❓")

    findings = report.get("findings", [])
    if findings:
        st.subheader("Verifier Results")
        for idx, f in enumerate(findings, start=1):
            fid = f.get("id")
            item = f.get("llm_item")
            verifier = f.get("verifier", {})
            status = verifier.get("status", "insufficient evidence")
            document_excerpt = f.get("document_excerpt", "No related document text found.")
            policy_excerpt = f.get("policy_excerpt", "No related policy text found.")
            citation = verifier.get("citation", "No citation")

            if status == "pass":
                status_tag = "✔️ Pass"
            elif status == "fail":
                status_tag = "❌ Fail"
            else:
                status_tag = "❓ Further review needed"

            with st.expander(f"Issue {idx}: {item} — {status_tag}", expanded=False):
                st.markdown(f"**Original text from User Document:**\n{document_excerpt}")
                st.markdown("**Verifier output:**")
                st.markdown(f"- Status: **{status_tag}**")
                st.markdown(f"- Rationale: {verifier.get('rationale')}")
                st.markdown(f"- Suggested edit: {verifier.get('suggested_edit')}")
                st.markdown(f"**Relevant Policy text:**\n{policy_excerpt}")
                st.markdown(f"**Citation:** {citation}")

        def build_word_report(report_data):
            doc = Document()
            doc.add_heading('Compliance Verifier Report', level=1)
            doc.add_paragraph(f"Policy version: {report_data.get('policy_version', 'n/a')}")
            doc.add_paragraph(f"Total findings: {summary.get('total_findings', 0)}")
            doc.add_paragraph(f"Pass: {pass_count}")
            doc.add_paragraph(f"Fail: {fail_count}")
            doc.add_paragraph('')

            for idx, f in enumerate(report_data.get('findings', []), start=1):
                verifier = f.get('verifier', {})
                doc.add_heading(f'Issue {idx}', level=2)
                doc.add_paragraph(f"LLM issue: {f.get('llm_item', '')}")
                doc.add_paragraph('Original text from User Document:')
                doc.add_paragraph(f.get('document_excerpt', 'No related document text found.'))
                doc.add_paragraph('Verifier output:')
                doc.add_paragraph(f"Status: {verifier.get('status', 'insufficient evidence')}")
                doc.add_paragraph(f"Rationale: {verifier.get('rationale', '')}")
                doc.add_paragraph(f"Suggested edit: {verifier.get('suggested_edit', '')}")
                doc.add_paragraph('Relevant Policy text:')
                doc.add_paragraph(f.get('policy_excerpt', 'No related policy text found.'))
                doc.add_paragraph(f"Citation: {verifier.get('citation', 'No citation')}")
                doc.add_paragraph('')

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()

        word_bytes = build_word_report(report)
        st.download_button(
            label="Download formatted results as Word document",
            data=word_bytes,
            file_name="compliance_verification_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.info("No findings were generated.")
else:
    st.info("Upload a policy PDF and a target document PDF, then run the verification workflow.")
